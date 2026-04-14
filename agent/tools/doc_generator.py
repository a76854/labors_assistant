import os
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from langchain_core.tools import tool

from agent.state import LawsuitElementsSchema


PROJECT_ROOT = Path(__file__).resolve().parents[2]


@tool(args_schema=LawsuitElementsSchema)
def generate_legal_doc_tool(
  plaintiff: str,
  defendant: str,
  claim: str,
  amount: str,
  cause_of_action: str,
  facts_and_reasons: str,
  court_name: str,
) -> str:
  """
  法律文书生成工具：生成符合中国法院常用格式的《民事起诉状》。

  '【终极动作】当且仅当：1.已集齐四个案件要素；2.已调用检索工具查阅了法律；3.已向用户输出了案情分析评估后。你必须作为最后一步调用此工具，为用户生成标准的法律文书。'

  参数:
    plaintiff: 原告信息。
    defendant: 被告信息。
    claim: 诉讼请求。
    amount: 涉案金额。
    cause_of_action: 案由。
    facts_and_reasons: 事实与理由。
    court_name: 管辖法院名称。

  返回:
    生成结果说明与模拟下载链接。
  """
  try:
    # 固定保存到项目根目录，避免因启动目录不同导致后端下载接口找不到文件。
    output_dir = PROJECT_ROOT / "generated_docs"
    output_dir.mkdir(parents=True, exist_ok=True)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"legal_doc_{timestamp}.docx"
    file_path = output_dir / filename

    doc = Document()

    # 统一正文样式，便于法院窗口打印审阅。
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "宋体"
    normal_style.font.size = Pt(12)

    title = doc.add_paragraph("民事起诉状")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.name = "宋体"
    title.runs[0].font.size = Pt(22)
    title.runs[0].bold = True
    doc.add_paragraph("")

    doc.add_paragraph(f"原告：{plaintiff}")
    doc.add_paragraph(f"被告：{defendant}")
    doc.add_paragraph("")

    cause_paragraph = doc.add_paragraph()
    cause_label = cause_paragraph.add_run("案由：")
    cause_label.bold = True
    cause_paragraph.add_run("\n")
    cause_paragraph.add_run(cause_of_action)

    claim_paragraph = doc.add_paragraph()
    claim_label = claim_paragraph.add_run("诉讼请求：")
    claim_label.bold = True
    claim_paragraph.add_run("\n")
    claim_paragraph.add_run(claim)
    claim_paragraph.add_run(f"\n诉讼标的金额：{amount}")
    doc.add_paragraph("")

    facts_paragraph = doc.add_paragraph()
    facts_label = facts_paragraph.add_run("事实与理由：")
    facts_label.bold = True
    facts_paragraph.add_run("\n")
    facts_paragraph.add_run(facts_and_reasons)
    facts_paragraph.paragraph_format.first_line_indent = Cm(0.74)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph("此致")
    doc.add_paragraph(court_name)

    doc.add_paragraph("")
    doc.add_paragraph("")

    signer_paragraph = doc.add_paragraph("起诉人：")
    signer_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_paragraph = doc.add_paragraph("年    月    日")
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(str(file_path))

    download_base_url = os.getenv("DOCUMENT_DOWNLOAD_BASE_URL", "http://127.0.0.1:8000").rstrip("/")
    download_link = f"{download_base_url}/api/v1/download/{filename}"
    return (
      "文书已生成成功。\n"
      f"文件名：{filename}\n"
      f"本地路径：{file_path.as_posix()}\n"
      f"下载链接：{download_link}\n"
      f"点击下载：[{filename}]({download_link})"
    )
  except Exception as exc:
    return f"文书生成失败：{exc}"
